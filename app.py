import os
from flask import Flask, render_template, request, redirect, url_for
from werkzeug.utils import secure_filename
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from azure.core.credentials import AzureKeyCredential
from azure.ai.translation.text import TextTranslationClient
from datetime import datetime, timedelta
import uuid
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF
import io

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# Load environment variables
translator_key = os.getenv('TRANSLATOR_KEY')
translator_region = os.getenv('TRANSLATOR_REGION')
translator_endpoint = os.getenv('TRANSLATOR_ENDPOINT')
blob_conn_str = os.getenv('BLOB_CONNECTION_STRING')
container_name = os.getenv('BLOB_CONTAINER_NAME', 'translated-docs')

# Initialize Azure clients
blob_service_client = BlobServiceClient.from_connection_string(blob_conn_str)
translator_credential = AzureKeyCredential(translator_key)
translator_client = TextTranslationClient(
    endpoint=translator_endpoint,
    credential=translator_credential,
    region=translator_region
)

# Ensure container exists
container_client = blob_service_client.get_container_client(container_name)
if not container_client.exists():
    container_client.create_container()

def extract_text(file_stream, filename):
    if filename.endswith('.docx'):
        doc = Document(file_stream)
        return '\n'.join([p.text for p in doc.paragraphs])
    elif filename.endswith('.pdf'):
        reader = PdfReader(file_stream)
        return ''.join([page.extract_text() for page in reader.pages])

def create_translated_file(text, filename):
    if filename.endswith('.docx'):
        doc = Document()
        doc.add_paragraph(text)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    elif filename.endswith('.pdf'):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, text)
        output = io.BytesIO()
        pdf.output(output)
        output.seek(0)
        return output

def generate_sas(blob_name):
    return generate_blob_sas(
        account_name=blob_service_client.account_name,
        container_name=container_name,
        blob_name=blob_name,
        account_key=blob_service_client.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=1)
    )

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/translate', methods=['POST'])
def translate():
    file = request.files['file']
    target_lang = request.form['target_lang']
    
    if not file or not allowed_file(file.filename):
        return "Invalid file", 400
    
    filename = secure_filename(file.filename)
    unique_id = uuid.uuid4().hex
    original_blob = f"{unique_id}/original_{filename}"
    translated_blob = f"{unique_id}/translated_{filename}"
    
    # Extract text
    file_stream = io.BytesIO(file.read())
    text = extract_text(file_stream, filename)
    
    # Translate
    response = translator_client.translate(content=[text], to=[target_lang])
    translated_text = response[0].translations[0].text
    
    # Create translated file
    translated_file = create_translated_file(translated_text, filename)
    
    # Upload to Blob
    file_stream.seek(0)
    container_client.upload_blob(original_blob, file_stream)
    translated_file.seek(0)
    container_client.upload_blob(translated_blob, translated_file)
    
    # Generate URLs
    base_url = f"https://{blob_service_client.account_name}.blob.core.windows.net"
    original_url = f"{base_url}/{container_name}/{original_blob}?{generate_sas(original_blob)}"
    translated_url = f"{base_url}/{container_name}/{translated_blob}?{generate_sas(translated_blob)}"
    
    return render_template('result.html', original=original_url, translated=translated_url)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx'}

if __name__ == '__main__':
    app.run()